from docx import *
import os
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement, ns
import datetime
import pandas
import json
import csv
import matplotlib.pyplot as plt
from docx.oxml.ns import qn

# variable of list page
list_pageNumber = 0
list_pagesubNumber = 0
list_pagesubNumber2 = 0

# variable of figure number
list_figure = 0
colorStrDefalt = '8ED680'


def autoStart(name):
    filename = name

    def create_element(name):
        return OxmlElement(name)

    def create_attribute(element, name, value):
        element.set(ns.qn(name), value)

    def add_page_number(run):
        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    # set repeat table row on every new page

    def set_repeat_table_header(row):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        return row

    # Make file CSV to JSON

    def makeJson(csvFilePath, jsonFilePath):
        data = {}
        with open(csvFilePath, encoding='utf-8') as csvf:
            csvReader = csv.DictReader(csvf)
            key_id = 0
            for rows in csvReader:
                key = key_id
                data[key] = rows
                key_id += 1
        with open(jsonFilePath, 'w', encoding='utf-8') as jsonf:
            jsonf.write(json.dumps(data, indent=4))

    DataJson = open(
        "D:/github/WebReport/backend/api/sources/dataFile.json", "w")
    DataJson.close()

    csvFilePath = r'D:/github/WebReport/backend/uploads/'+name
    jsonFilePath = r'D:/github/WebReport/backend/api/sources/dataFile.json'

    makeJson(csvFilePath, jsonFilePath)

    # create variable global
    DataHost = {}
    CountID = {}
    GraphAll = {}
    GraphTotol = {}
    Risk = {}
    Protocol = {}
    dataALL = {}
    datalist = []

    vulnerability = {}
    vulnelist = []
    name = {}
    hosts = ""
    ports = ""
    portlist = []

    level4 = []
    level3 = []
    level2 = []
    level1 = []
    count = 0
    num = 0
    count_table = 0
    host = ""
    port = ""

    DataJSON = pandas.read_json(jsonFilePath)
    DataCSV = pandas.read_csv(csvFilePath)

    PluginID = [DataJSON[i]["Host"]
                for i in DataJSON if DataJSON[i]["Risk"] != "None"]

    for i in PluginID:
        if i in CountID:
            CountID[i] += 1
        else:
            CountID[i] = 1

    for i in PluginID:
        for j in DataJSON:
            if DataJSON[j]["Risk"] != 'None':
                if DataJSON[j]["Host"] == i:
                    if DataJSON[j]["Risk"] in Risk:
                        Risk[DataJSON[j]["Risk"]] += 1
                    else:
                        Risk[DataJSON[j]["Risk"]] = 1
                    count += 1
                    num += 1
                    if count == CountID[i]:
                        GraphAll[i] = Risk
                        Risk = {}
                        count = 0
                        if num == (len(PluginID) - 1):
                            break

    Critical = 0
    High = 0
    Medium = 0
    Low = 0

    for i in GraphAll:
        if "Critical" in GraphAll[i]:
            GraphTotol["Critical"] = int(GraphAll[i]["Critical"])
            Critical += int(GraphAll[i]["Critical"])

        if "High" in GraphAll[i]:
            GraphTotol["High"] = int(GraphAll[i]["High"])
            High += int(GraphAll[i]["High"])

        if "Medium" in GraphAll[i]:
            GraphTotol["Medium"] = int(GraphAll[i]["Medium"])
            Medium += int(GraphAll[i]["Medium"])

        if "Low" in GraphAll[i]:
            GraphTotol["Low"] = int(GraphAll[i]["Low"])
            Low += int(GraphAll[i]["Low"])

        GraphTotol["Critical"] = Critical
        GraphTotol["High"] = High
        GraphTotol["Medium"] = Medium
        GraphTotol["Low"] = Low

    names = 'Critical', 'High', 'Medium', 'Low'
    values = [GraphTotol["Critical"], GraphTotol["High"],
              GraphTotol["Medium"], GraphTotol["Low"]]

    plt.pie(values, labels=names, labeldistance=1.15)
    plt.legend(loc='upper center', bbox_to_anchor=(0.5, 0),
               fancybox=True, shadow=True, ncol=4)
    plt.savefig("D:/github/WebReport/backend/api/image/Graph.png")

    for i in DataJSON:
        if DataJSON[i]["Name"] == "Nessus SYN scanner":
            if DataJSON[i]["Host"] in DataHost:
                DataHost[DataJSON[i]["Host"]] += 1
            else:
                DataHost[DataJSON[i]["Host"]] = 1

    for i in DataJSON:

        if DataJSON[i]["Name"] == "Nessus SYN scanner":
            host = DataJSON[i]["Host"]
            port = port + ',' + DataJSON[i]["Port"]
            protocol = DataJSON[i]["Protocol"]
            count_table += 1

            if count_table == DataHost[host]:
                dataALL["Host"] = host
                dataALL["Port"] = protocol.upper() + "" + port
                datalist.append(dataALL)
                port = ""
                count_table = 0
                dataALL = {}

    for i in DataJSON:
        if DataJSON[i]["Risk"] != "None":
            if DataJSON[i]["Name"] in name:
                name[DataJSON[i]["Name"]] += 1
            else:
                name[DataJSON[i]["Name"]] = 1

    for i in DataJSON:
        count = 0
        if DataJSON[i]["Risk"] != "None":
            vulnerability["Name"] = DataJSON[i]["Name"]
            vulnerability["Synopsis"] = DataJSON[i]["Synopsis"]
            vulnerability["Description"] = DataJSON[i]["Description"]
            vulnerability["Solution"] = DataJSON[i]["Solution"]
            vulnerability["Remark"] = DataJSON[i]["See Also"]
            vulnerability["Protocol"] = DataJSON[i]["Protocol"]
            if DataJSON[i]["Risk"] == "Critical":
                vulnerability["Risk"] = 4
            if DataJSON[i]["Risk"] == "High":
                vulnerability["Risk"] = 3
            if DataJSON[i]["Risk"] == "Medium":
                vulnerability["Risk"] = 2
            if DataJSON[i]["Risk"] == "Low":
                vulnerability["Risk"] = 1
            for j in DataJSON:
                if vulnerability["Name"] == DataJSON[j]["Name"]:
                    count += 1
                    if hosts == "":
                        hosts = DataJSON[j]["Host"]
                    else:
                        hosts = hosts + "," + DataJSON[j]["Host"]

                    portlist.append(DataJSON[j]["Port"])
                    if count == name[vulnerability["Name"]]:
                        portlist = list(dict.fromkeys(portlist))
                        for k in portlist:
                            if ports == "":
                                ports = k
                            else:
                                ports = ports + "," + k
                        ports = vulnerability["Protocol"].upper() + " " + ports
                        vulnerability["Host"] = hosts
                        vulnerability["Port"] = ports
                        vulnelist.append(vulnerability)
                        count = 0
                        vulnerability = {}
                        hosts = ""
                        ports = ""
                        portlist = []
                        break

    vulnelist = [dict(i) for i in {tuple(j.items()) for j in vulnelist}]

    for i in vulnelist:
        if i["Risk"] == 4:
            level4.append(i)
        if i["Risk"] == 3:
            level3.append(i)
        if i["Risk"] == 2:
            level2.append(i)
        if i["Risk"] == 1:
            level1.append(i)

    levelall = level4 + level3 + level2 + level1

    doc = Document()

    def tabBgColor(table, cols, cols2, rows, colorStr):
        shading_list = locals()
        i = cols
        cols2 = cols2
        row = rows
        shading_list['shading_elm_' + str(i)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStr))
        table.rows[row].cells[cols2]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(i)])

    def autoTable(id, data):  # Creating a table object
        global colorStrDefalt
        table = doc.add_table(rows=6, cols=4, style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.allow_autofit = False
        for i in range(3):
            table.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for j in range(6):
                table.cell(j, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        shading_list = locals()
        shading_list['shading_elm_' + str(4)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStrDefalt))
        table.rows[0].cells[0]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(4)])

        shading_list = locals()
        shading_list['shading_elm_' + str(4)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStrDefalt))
        table.rows[1].cells[0]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(4)])

        shading_list = locals()
        shading_list['shading_elm_' + str(4)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStrDefalt))
        table.rows[2].cells[0]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(4)])

        shading_list = locals()
        shading_list['shading_elm_' + str(4)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStrDefalt))
        table.rows[3].cells[0]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(4)])

        shading_list = locals()
        shading_list['shading_elm_' + str(4)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStrDefalt))
        table.rows[4].cells[0]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(4)])

        shading_list = locals()
        shading_list['shading_elm_' + str(4)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStrDefalt))
        table.rows[5].cells[0]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(4)])

        shading_list = locals()
        shading_list['shading_elm_' + str(4)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStrDefalt))
        table.rows[0].cells[2]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(4)])

        shading_list = locals()
        shading_list['shading_elm_' + str(4)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStrDefalt))
        table.rows[1].cells[2]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(4)])

        # Header Table
        nametable800 = table.rows[0].cells[2]
        nametable800.text = "Finding"
        rowalignment = nametable800.paragraphs[0]
        table.rows[0].cells[2].paragraphs[0].runs[0].font.bold = True
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER

        nametable801 = table.rows[0].cells[3]
        nametable801.text = str(data["Name"])
        rowalignment = nametable801.paragraphs[0]
        # rowalignment.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        nametable802 = table.rows[1].cells[2]
        nametable802.text = "Port"
        rowalignment = nametable802.paragraphs[0]
        table.rows[1].cells[2].paragraphs[0].runs[0].font.bold = True
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER

        nametable803 = table.rows[1].cells[3]
        nametable803.text = str(data["Port"])
        rowalignment = nametable803.paragraphs[0]
        # rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Define row in table 2D example row10 is meant row index 1 colume index 0
        row00 = table.rows[0].cells[0]
        row00.text = "ID"
        rowalignment = row00.paragraphs[0]
        # rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True

        row01 = table.rows[0].cells[1]
        row01.text = str(id)
        rowalignment = row01.paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER

        row10 = table.rows[1].cells[0]
        row10.text = "Severity"
        rowalignment = row10.paragraphs[0]
        # rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True

        row11 = table.rows[1].cells[1]
        colsNum = 4
        colsNum2 = 1
        row = 1
        if data["Risk"] == 4:
            row11.text = "Critical"
            colorStr = 'C600EE'
            tabBgColor(table, colsNum, colsNum2, row, colorStr)
        if data["Risk"] == 3:
            row11.text = "High"
            colorStr = 'FF0000'
            tabBgColor(table, colsNum, colsNum2, row, colorStr)
        if data["Risk"] == 2:
            row11.text = "Medium"
            colorStr = 'FFCD00'
            tabBgColor(table, colsNum, colsNum2, row, colorStr)
        if data["Risk"] == 1:
            row11.text = "Low"
            colorStr = 'EFFF00'
            tabBgColor(table, colsNum, colsNum2, row, colorStr)
        rowalignment = row11.paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[1].cells[1].paragraphs[0].runs[0].font.bold = True

        row20 = table.rows[2].cells[0]
        row20.text = "Target"
        rowalignment = row20.paragraphs[0]
        table.rows[2].cells[0].paragraphs[0].runs[0].font.bold = True
        # rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row21 = table.rows[2].cells[1]
        row21.text = data["Host"]
        rowalignment = row21.paragraphs[0]
        # rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER

        row30 = table.rows[3].cells[0]
        row30.text = "Detail"
        rowalignment = row30.paragraphs[0]
        table.rows[3].cells[0].paragraphs[0].runs[0].font.bold = True
        # rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row31 = table.rows[3].cells[1]
        row31.text = data["Description"]
        # rowalignment = row31.paragraphs[0]
        # rowalignment.alignment = WD_ALIGN_PARAGRAPH.LEFT

        row40 = table.rows[4].cells[0]
        row40.text = "Solution"
        rowalignment = row40.paragraphs[0]
        table.rows[4].cells[0].paragraphs[0].runs[0].font.bold = True
        # rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row41 = table.rows[4].cells[1]
        row41.text = data["Solution"]
        # rowalignment = row31.paragraphs[0]
        # rowalignment.alignment = WD_ALIGN_PARAGRAPH.LEFT

        row50 = table.rows[5].cells[0]
        row50.text = "Remark"
        rowalignment = row50.paragraphs[0]
        table.rows[5].cells[0].paragraphs[0].runs[0].font.bold = True
        # rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row51 = table.rows[5].cells[1]
        row51.text = data["Remark"]
        # rowalignment = row31.paragraphs[0]
        # rowalignment.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # A merge is specified
        col21 = table.cell(2, 1)
        col23 = table.cell(2, 3)
        col21.merge(col23)
        col31 = table.cell(3, 1)
        col33 = table.cell(3, 3)
        col31.merge(col33)
        col41 = table.cell(4, 1)
        col43 = table.cell(4, 3)
        col41.merge(col43)
        col51 = table.cell(5, 1)
        col53 = table.cell(5, 3)
        col51.merge(col53)

        for cell in table.columns[0].cells:
            cell.width = Inches(0.8)
        for cell in table.columns[1].cells:
            cell.width = Inches(0.85)
        for cell in table.columns[2].cells:
            cell.width = Inches(0.8)
        for cell in table.columns[3].cells:
            cell.width = Inches(3.8)

    # page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Define Variable globel
    company = "บริษัท ไอเน็ต แมเนจด์ เซอร์วิสเซส จำกัด"

    # Set page size A4
    doc.sections[0].page_height = Cm(29.7)
    doc.sections[0].page_width = Cm(21)

    # Set Font Page
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(16)

    # Define header image
    section = doc.sections[0]
    header = section.header_distance = Inches(0)
    # header = section.right_margin = Inches(0)
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = 2
    header_run = paragraph.add_run()
    header_run = header_run.add_picture(
        'D:/github/WebReport/backend/api/image/inetms_logo.png', width=Inches(2.2), height=Inches(1.5))
    paragraph.paragraph_format.right_indent = -Inches(1.25)

    FooterNumber = doc.sections[0].footer
    FooterNumber = FooterNumber.paragraphs[0]
    FooterNumber.alignment = 2
    FooterNumber = FooterNumber.add_run(
        "(Private and Confidential, Do not release to publicity except allow from {})   Page  ".format(company))
    FooterNumber.font.size = Pt(11)
    add_page_number(FooterNumber)

    # space 3 paragraph
    doc.add_heading()
    doc.add_heading()
    doc.add_heading()

    # Define Font size for Header
    header = doc.add_paragraph()
    header.alignment = 2
    header = header.add_run("Vulnerability Assessment Report \n For")
    header.font.size = Pt(28)
    header.bold = True

    # Add Picture To Header
    logo = doc.add_picture('D:/github/WebReport/backend/api/image/inetms_logo2.png',
                           width=Inches(6), height=Inches(2))
    logo = doc.paragraphs[-1]
    logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Define name company
    header = doc.add_paragraph()
    header.alignment = 2
    header = header.add_run(company)
    header.font.size = Pt(18)
    header.bold = True

    # Define Font size for content
    DateNow = doc.add_paragraph()
    DateNow.alignment = 2
    DateFormat = datetime.datetime.now()
    DateFormat = str(DateFormat.strftime("%d")) + " " + str(DateFormat.strftime("%B")) + " " + str(
        DateFormat.strftime("%Y"))
    DateNow = DateNow.add_run("Date : " + DateFormat)
    DateNow.font.size = Pt(18)
    DateNow.bold = True

    # Adding a page break
    doc.add_page_break()

    #################################################################   Next page   ############################################################################

    # Define content
    contentPage2 = doc.add_paragraph()
    contentPage2 = contentPage2.add_run(
        "Document Security Level : Confidential \nDocument Version: 1.0")
    contentPage2.font.size = Pt(16)
    contentPage2.bold = True

    # Data For test table1
    data = [
        {'version': '1.0', 'date': 'DateFormat',
         'editreport': 'Creation', 'editor': 'INET Managed Service'}
    ]

    # Creating a table1 object
    table1 = doc.add_table(rows=1, cols=4, style='Table Grid')
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(4):
        table1.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i in range(4):
        shading_list = locals()
        shading_list['shading_elm_' + str(4)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStrDefalt))
        table1.rows[0].cells[i]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(4)])

    def make_rows_bold(*rows):
        for row in rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    row = table1.rows[0].cells
    row[0].text = "Version"
    row[1].text = "Date"
    row[2].text = "Edit Report"
    row[3].text = "Editor"

    make_rows_bold(table1.rows[0])
    set_repeat_table_header(table1.rows[0])

    table1col0 = row[0].paragraphs[0]
    table1col1 = row[1].paragraphs[0]
    table1col2 = row[2].paragraphs[0]
    table1col3 = row[3].paragraphs[0]
    table1col0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table1col1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table1col2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table1col3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(len(data)):
        row = table1.add_row().cells
        row[0].text = data[i]['version']
        rowalignment = row[0].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[1].text = str(DateFormat)
        rowalignment = row[1].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[2].text = data[i]['editreport']
        rowalignment = row[2].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[3].text = data[i]['editor']
        rowalignment = row[3].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cell in table1.columns[0].cells:
        cell.width = Inches(0.8)
    for cell in table1.columns[3].cells:
        cell.width = Inches(2.3)

    # space 1 paragraph
    doc.add_paragraph()

    # Creating a table2 object
    table2 = doc.add_table(rows=5, cols=4, style='Table Grid')
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(4):
        table2.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i in range(4):
        shading_list = locals()
        shading_list['shading_elm_' + str(4)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStrDefalt))
        table2.rows[0].cells[i]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(4)])

    # Header Table
    nametable00 = table2.rows[0].cells[0]
    nametable02 = table2.rows[0].cells[2]

    set_repeat_table_header(table2.rows[0])

    row10 = table2.rows[1].cells[0]
    row20 = table2.rows[2].cells[0]
    row30 = table2.rows[3].cells[0]
    row40 = table2.rows[4].cells[0]

    row12 = table2.rows[1].cells[2]
    row22 = table2.rows[2].cells[2]
    row32 = table2.rows[3].cells[2]
    row42 = table2.rows[4].cells[2]

    # A merge is specified
    col0 = table2.cell(0, 0)
    col1 = table2.cell(0, 1)
    col0.merge(col1)
    col2 = table2.cell(0, 2)
    col3 = table2.cell(0, 3)
    col2.merge(col3)

    # Define Topic in table
    nametable00.text = "Example Company Limited."
    nametable02.text = "INET Managed Services CO., LTD."

    make_rows_bold(table2.rows[0])

    rowalignment = nametable00.paragraphs[0]
    rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rowalignment = nametable02.paragraphs[0]
    rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Define Topic in Row
    row10.text = "Name"

    row20.text = "Position"

    row30.text = "Tel"

    row40.text = "Signature"

    row12.text = "Name"

    row22.text = "Position"

    row32.text = "Tel"

    row42.text = "Signature"

    for i in range(5):
        table2.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        table2.rows[i].cells[2].paragraphs[0].runs[0].font.bold = True

    rowalignment = row10.paragraphs[0]
    rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rowalignment = row20.paragraphs[0]
    rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rowalignment = row30.paragraphs[0]
    rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rowalignment = row40.paragraphs[0]
    rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rowalignment = row12.paragraphs[0]
    rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rowalignment = row22.paragraphs[0]
    rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rowalignment = row32.paragraphs[0]
    rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rowalignment = row42.paragraphs[0]
    rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Define Color in table
    table2.style = 'Table Grid'

    for cell in table2.columns[0].cells:
        cell.width = Inches(0.85)
    for cell in table2.columns[1].cells:
        cell.width = Inches(2.02)
    for cell in table2.columns[2].cells:
        cell.width = Inches(0.85)
    for cell in table2.columns[3].cells:
        cell.width = Inches(2.51)

    # Break Page
    doc.add_page_break()

    # Header page 3
    header3 = doc.add_paragraph()
    header3.alignment = 0
    header3 = header3.add_run("Table of Contents")  # company is variable top
    header3.font.size = Pt(16)
    header3.font.color.rgb = RGBColor(1, 0, 254)
    header3.bold = True

    # Code for making Table of Contents

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    # change 1-3 depending on heading levels you need
    instrText.text = 'TOC \o "1-3" \h \z'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p
    # Break Page
    doc.add_page_break()

    # Make funtion for list number figure

    def figure_number(title):
        global list_figure
        figureNumber = doc.add_paragraph()
        list_figure += 1
        figureNumber.alignment = 1
        data = "Figure " + str(list_figure) + ':  ' + title
        figureNumber.add_run(data)

    ########################### Page 4  ######################################

    # Header page 4
    heading = doc.add_heading(
        "1.  Restrictions on disclosure and use of information", 1)
    title_style = heading.style
    title_style.font.size = Pt(16)
    rFonts = title_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), 'TH Sarabun New')

    # doc.add_paragraph()

    paragraph = doc.add_paragraph(
        "     Restriction on Disclosure and Use of Confidential Information. The Executive understands and agrees that the Confidential Information constitutes an asset of the Company and its affiliated entities and may not be converted to the Executive's own use. Accordingly, the Executive hereby agrees that the Executive shall not, directly, or indirectly, at any time, reveal, divulge, or disclose to any Person not expressly authorized by the Company any Confidential Information, and the Executive shall not, directly, or indirectly, use or make use of any Confidential Information in connection with any business activity other than that of the Company. The parties acknowledge and agree that this Agreement is not intended to, and does not, alter either the Company's rights or the Executive's obligations under any state or federal statutory or common law regarding trade secrets and unfair trade.")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Break Page 4
    doc.add_page_break()

    ########################### Page 5  ######################################

    # Header page 5
    doc.add_heading("2. Operation Method", 1)
    heading = doc.add_heading("   2.1 Posture Review", 2)
    doc.add_heading("   2.2 Information Gathering", 2)
    doc.add_heading("   2.3 Enumeration", 2)
    doc.add_heading("   2.4 Vulnerability Assessment", 2)
    doc.add_heading("   2.5 Analyze & Evaluate Risk Value", 2)
    doc.add_heading("   2.6 Report", 2)
    title_style = heading.style
    title_style.font.size = Pt(16)
    rFonts = title_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), 'TH Sarabun New')
    # doc.add_paragraph()

    # Add Figure 1 page 5
    Figure1 = doc.add_picture('D:/github/WebReport/backend/api/image/operation method.jpg',
                              width=Cm(14.3), height=Cm(7))
    Figure1 = doc.paragraphs[-1]
    Figure1.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Define image to right

    figure_number("Operation Method")

    # Break Page 5
    doc.add_page_break()

    ########################### Page 6  ######################################
    sections = doc.sections

    # Page dimensions Landscape
    section6 = doc.add_section(WD_SECTION.ODD_PAGE)
    section6.orientation = WD_ORIENT.LANDSCAPE
    section6.page_width = Cm(29.7)
    section6.page_height = Cm(21)

    # Header page 6
    doc.add_heading("3. Project Scope", 1)
    doc.add_heading("   3.1 Infrastructure Vulnerability Assessment", 2)
    # doc.add_paragraph()

    # Title table 3
    table3Title = doc.add_paragraph()
    table3Title = table3Title.add_run(
        "       Target / IP Address:")  # company is variable top
    table3Title.font.size = Pt(16)
    table3Title.font.color.rgb = RGBColor(2, 0, 0)
    table3Title.bold = True

    # Create table of target / IP ADDRESS:
    table3 = doc.add_table(rows=1, cols=8, style='Table Grid')
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i in range(8):
        table3.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i in range(8):
        shading_list = locals()
        shading_list['shading_elm_' + str(8)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStrDefalt))
        table3.rows[0].cells[i]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(8)])

    row = table3.rows[0].cells
    row[0].text = "No"
    row[1].text = "Domain / Server Name"
    row[2].text = "Public IP Address"
    row[3].text = "Private IP Address"
    row[4].text = "OS/Model"
    row[5].text = "Functions"
    row[6].text = "Public Assessment"
    row[7].text = "Private Assessment"

    make_rows_bold(table3.rows[0])
    set_repeat_table_header(table3.rows[0])

    table3col0 = row[0].paragraphs[0]
    table3col0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table3col1 = row[1].paragraphs[0]
    table3col1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table3col2 = row[2].paragraphs[0]
    table3col2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table3col3 = row[3].paragraphs[0]
    table3col3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table3col4 = row[4].paragraphs[0]
    table3col4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table3col5 = row[5].paragraphs[0]
    table3col5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table3col6 = row[6].paragraphs[0]
    table3col6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table3col7 = row[7].paragraphs[0]
    table3col7.alignment = WD_ALIGN_PARAGRAPH.CENTER

    numberoftable = 0
    for i in datalist:
        row = table3.add_row().cells
        numberoftable += 1
        row[0].text = str(numberoftable)
        rowalignment = row[0].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER

        row[1].text = "None"
        rowalignment = row[1].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER

        row[2].text = "None"
        rowalignment = row[2].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER

        row[3].text = i["Host"]
        rowalignment = row[3].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.LEFT

        row[4].text = "None"
        rowalignment = row[4].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER

        row[5].text = "None"
        rowalignment = row[5].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER

        row[6].text = "None"
        rowalignment = row[6].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER

        row[7].text = "None"
        rowalignment = row[7].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cell in table3.columns[0].cells:
        cell.width = Inches(0.49)

    # Break Page
    doc.add_page_break()

    doc.add_heading("   3.2 Web Application Vulnerability Assessment", 2)
    # doc.add_paragraph()

    # Title table 4
    table3Title = doc.add_paragraph()
    table3Title = table3Title.add_run(
        "       Target / IP Address:")  # company is variable top
    table3Title.font.size = Pt(16)
    table3Title.font.color.rgb = RGBColor(2, 0, 0)
    table3Title.bold = True

    # Create table of target / IP ADDRESS:
    table4 = doc.add_table(rows=2, cols=8, style='Table Grid')
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(8):
        table4.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i in range(8):
        shading_list = locals()
        shading_list['shading_elm_' + str(8)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStrDefalt))
        table4.rows[0].cells[i]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(8)])

    row = table4.rows[0].cells
    row[0].text = "No"
    row[1].text = "Domain / Server Name"
    row[2].text = "Public IP Address"
    row[3].text = "Private IP Address"
    row[4].text = "OS/Model"
    row[5].text = "Functions"
    row[6].text = "Public Assessment"
    row[7].text = "Private Assessment"

    make_rows_bold(table4.rows[0])
    set_repeat_table_header(table4.rows[0])

    table4col0 = row[0].paragraphs[0]
    table4col0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table4col1 = row[1].paragraphs[0]
    table4col1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table4col2 = row[2].paragraphs[0]
    table4col2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table4col3 = row[3].paragraphs[0]
    table4col3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table4col4 = row[4].paragraphs[0]
    table4col4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table4col5 = row[5].paragraphs[0]
    table4col5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table4col6 = row[6].paragraphs[0]
    table4col6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table4col7 = row[7].paragraphs[0]
    table4col7.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cell in table4.columns[0].cells:
        cell.width = Inches(0.4)

    # Break Page 6
    doc.add_page_break()

    ########################### Page 9  ######################################

    # Page dimensions Landscape
    section6 = doc.add_section(WD_SECTION.ODD_PAGE)
    section6.orientation = WD_ORIENT.PORTRAIT
    section6.page_width = Cm(21)
    section6.page_height = Cm(29.7)

    doc.add_heading("4. Testing Tools", 1)
    # doc.add_paragraph()

    # Make table5 Testing Tools
    table5 = doc.add_table(rows=1, cols=2, style='Table Grid')
    table5.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(2):
        table5.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i in range(2):
        shading_list = locals()
        shading_list['shading_elm_' + str(2)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStrDefalt))
        table5.rows[0].cells[i]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(2)])

    row = table5.rows[0].cells
    row[0].text = "Tool Name"
    row[1].text = "Testing Type"

    make_rows_bold(table5.rows[0])
    set_repeat_table_header(table5.rows[0])

    table5col0 = row[0].paragraphs[0]
    table5col0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table5col1 = row[1].paragraphs[0]
    table5col1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    dataTool = [
        {"name": "Nmap", "type": "Host and Service Discovery"},
        {"name": "Nessus", "type": "Infrastructure Vulnerability Assessment"},
        {"name": "Acunetix", "type": "Web Application Vulnerability Assessment"},
    ]

    for i in range(len(dataTool)):
        datarow = table5.add_row().cells
        datarow[0].text = dataTool[i]["name"]
        row = datarow[0].paragraphs[0]
        row.alignment = WD_ALIGN_PARAGRAPH.CENTER
        datarow[1].text = dataTool[i]["type"]
        row = datarow[1].paragraphs[0]
        row.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cell in table5.columns[0].cells:
        cell.width = Cm(2.7)
    for cell in table5.columns[1].cells:
        cell.width = Cm(14)

    doc.add_heading("5. Infrastructure Vulnerability Assessment", 1)

    spacing = doc.add_paragraph()

    spacing1 = spacing.add_run(
        "    Vulnerability Assessment from Public Access \n")
    spacing2 = spacing.add_run("    Testing data :".format("exam") + "\n")
    spacing3 = spacing.add_run("    Tester IP Address :".format("exam"))

    spacing1.font.color.rgb = RGBColor(2, 0, 0)
    spacing1.bold = True
    spacing1.font.size = Pt(14)

    spacing2.font.color.rgb = RGBColor(2, 0, 0)
    spacing2.bold = True
    spacing2.font.size = Pt(14)

    spacing3.font.color.rgb = RGBColor(2, 0, 0)
    spacing3.bold = True
    spacing3.font.size = Pt(14)

    # Add Figure 2 page 7
    Figure2 = doc.add_picture(
        'D:/github/WebReport/backend/api/image/vulnerability public access.jpg', width=Cm(13), height=Cm(2.5))
    Figure2 = doc.paragraphs[-1]
    Figure2.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Define image to right

    figure_number("Vulnerability Assessment from Public Access")
    ########################### END Figure 2  ######################################

    spacing1 = doc.add_paragraph()
    spacing4 = spacing1.add_run(
        "    Vulnerability Assessment from Private Access (for private or restricted access target) \n")
    spacing5 = spacing1.add_run("    Testing data :".format("exam") + "\n")
    spacing6 = spacing1.add_run("    Tester IP Address :".format("exam"))

    spacing4.font.color.rgb = RGBColor(2, 0, 0)
    spacing4.bold = True
    spacing4.font.size = Pt(14)

    spacing5.font.color.rgb = RGBColor(2, 0, 0)
    spacing5.bold = True
    spacing5.font.size = Pt(14)

    spacing6.font.color.rgb = RGBColor(2, 0, 0)
    spacing6.bold = True
    spacing6.font.size = Pt(14)

    # Add Figure 8 page 10
    Figure3 = doc.add_picture(
        'D:/github/WebReport/backend/api/image/vulnerability public access.jpg', width=Cm(13), height=Cm(2.5))
    Figure3 = doc.paragraphs[-1]
    Figure3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    figure_number("Vulnerability Assessment from Private Access")
    # doc.add_paragraph()

    ########################### Page 11  ######################################

    # Page dimensions Landscape
    section8 = doc.add_section(WD_SECTION.ODD_PAGE)
    section8.orientation = WD_ORIENT.LANDSCAPE
    section8.page_width = Cm(29.7)
    section8.page_height = Cm(21)

    # Header page
    doc.add_heading("   5.1 Target Information", 2)
    # doc.add_paragraph()

    # Create table of target / IP ADDRESS:
    table6 = doc.add_table(rows=1, cols=5, style='Table Grid')
    table6.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(5):
        table6.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i in range(5):
        shading_list = locals()
        shading_list['shading_elm_' + str(5)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStrDefalt))
        table6.rows[0].cells[i]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(5)])

    row = table6.rows[0].cells
    row[0].text = "No"
    row[1].text = "Domain / Server Name"
    row[2].text = "IP Address"
    row[3].text = "OS/Model"
    row[4].text = "Port"
    make_rows_bold(table6.rows[0])
    set_repeat_table_header(table6.rows[0])

    table6col0 = row[0].paragraphs[0]
    table6col0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table6col1 = row[1].paragraphs[0]
    table6col1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table6col2 = row[2].paragraphs[0]
    table6col2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table6col3 = row[3].paragraphs[0]
    table6col3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table6col4 = row[4].paragraphs[0]
    table6col4.alignment = WD_ALIGN_PARAGRAPH.CENTER

    numberoftable = 0

    for i in datalist:
        row = table6.add_row().cells
        numberoftable += 1
        row[0].text = str(numberoftable)
        rowalignment = row[0].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[1].text = "None"
        rowalignment = row[1].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[2].text = i["Host"]
        rowalignment = row[2].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.LEFT
        row[3].text = "None"
        rowalignment = row[3].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[4].text = i["Port"]
        rowalignment = row[4].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for cell in table6.columns[0].cells:
        cell.width = Inches(0.5)
    for cell in table6.columns[2].cells:
        cell.width = Inches(1.2)
    for cell in table6.columns[3].cells:
        cell.width = Inches(1.3)
    for cell in table6.columns[4].cells:
        cell.width = Inches(4.5)

    # Break Page 11
    doc.add_page_break()

    ########################### Page 12  ######################################

    # Page dimensions Landscape
    section6 = doc.add_section(WD_SECTION.ODD_PAGE)
    section6.orientation = WD_ORIENT.PORTRAIT
    section6.page_width = Cm(21)
    section6.page_height = Cm(29.7)

    # Header page
    doc.add_heading("5.2 Executive summary", 2)
    # doc.add_paragraph()

    doc.add_paragraph(
        "         The purpose of this activity is to find the vulnerability on the target infrastructure.")

    # Header page
    heading = doc.add_heading(
        "      5.2.1 Summary Vulnerability by Severity", 3)
    title_style = heading.style
    title_style.font.size = Pt(16)
    rFonts = title_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), 'TH Sarabun New')
    doc.add_paragraph()

    # Add Figure 4 page 12
    Figure4 = doc.add_picture(
        'D:/github/WebReport/backend/api/image/Graph.png', width=Cm(11), height=Cm(8))
    Figure4 = doc.paragraphs[-1]
    Figure4.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Define image to right

    figure_number(
        "Summary by Severity of Infrastructure Vulnerability Assessment")
    doc.add_page_break()

    doc.add_heading("5.2.2 Vulnerability by Target", 3)
    # doc.add_paragraph()

    # Create table of Vulnerability by Target
    table7 = doc.add_table(rows=1, cols=8, style='Table Grid')
    table7.alignment = WD_TABLE_ALIGNMENT.CENTER
    colsNum = 8

    row = table7.rows[0].cells

    cell_00 = table7.cell(0, 0)
    cell_00.text = "No"

    cell_01 = table7.cell(0, 1)
    cell_01.text = "Domain/Server Name"

    cell_00 = table7.cell(0, 2)
    cell_00.text = "IP Address"

    for i in range(3):
        shading_list = locals()
        shading_list['shading_elm_' + str(colsNum)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStrDefalt))
        table7.rows[0].cells[i]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(colsNum)])

    cell_01 = table7.cell(0, 3)
    cell_01.text = "Critical"
    colsNum2 = 3
    rowx = 0
    colorStr = '8139B8'
    shading_list = locals()
    shading_list['shading_elm_' + str(colsNum)] = parse_xml(
        r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStr))
    table7.rows[rowx].cells[colsNum2]._tc.get_or_add_tcPr().append(
        shading_list['shading_elm_' + str(colsNum)])

    cell_00 = table7.cell(0, 4)
    cell_00.text = "High"
    colsNum2 = 4
    rowx = 0
    colorStr = 'FF0000'
    shading_list['shading_elm_' + str(colsNum)] = parse_xml(
        r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStr))
    table7.rows[rowx].cells[colsNum2]._tc.get_or_add_tcPr().append(
        shading_list['shading_elm_' + str(colsNum)])

    cell_01 = table7.cell(0, 5)
    cell_01.text = "Medium"
    colsNum2 = 5
    colorStr = 'FFD400'
    shading_list['shading_elm_' + str(colsNum)] = parse_xml(
        r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStr))
    table7.rows[rowx].cells[colsNum2]._tc.get_or_add_tcPr().append(
        shading_list['shading_elm_' + str(colsNum)])

    cell_00 = table7.cell(0, 6)
    cell_00.text = "Low"
    colsNum2 = 6
    colorStr = 'EFFF00'
    shading_list['shading_elm_' + str(colsNum)] = parse_xml(
        r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStr))
    table7.rows[rowx].cells[colsNum2]._tc.get_or_add_tcPr().append(
        shading_list['shading_elm_' + str(colsNum)])

    cell_01 = table7.cell(0, 7)
    cell_01.text = "Total"
    colsNum2 = 7
    colorStr = '0099FF'
    shading_list['shading_elm_' + str(colsNum)] = parse_xml(
        r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStr))
    table7.rows[rowx].cells[colsNum2]._tc.get_or_add_tcPr().append(
        shading_list['shading_elm_' + str(colsNum)])

    make_rows_bold(table7.rows[0])
    set_repeat_table_header(table7.rows[0])

    table7col0 = row[0].paragraphs[0]
    table7col0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table7col1 = row[1].paragraphs[0]
    table7col1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table7col2 = row[2].paragraphs[0]
    table7col2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table7col3 = row[3].paragraphs[0]
    table7col3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table7col4 = row[4].paragraphs[0]
    table7col4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table7col5 = row[5].paragraphs[0]
    table7col5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table7col6 = row[6].paragraphs[0]
    table7col6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table7col7 = row[7].paragraphs[0]
    table7col7.alignment = WD_ALIGN_PARAGRAPH.CENTER

    numberoftable = 0
    sum = 0
    Critical = 0
    High = 0
    Medium = 0
    Low = 0

    for i in range(8):
        table7.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i in GraphAll:
        row = table7.add_row().cells
        numberoftable += 1
        row[0].text = str(numberoftable)
        rowalignment = row[0].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[1].text = "None"
        rowalignment = row[1].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[2].text = i
        Critical = "0"
        High = "0"
        Medium = "0"
        Low = "0"

        if "Critical" in GraphAll[i]:
            Critical = int(GraphAll[i]["Critical"])
            row[3].text = str(Critical)
            rowalignment = row[3].paragraphs[0]
            rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[3].text = str(Critical)
        rowalignment = row[3].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if "High" in GraphAll[i]:
            High = int(GraphAll[i]["High"])
            row[4].text = str(High)
            rowalignment = row[4].paragraphs[0]
            rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[4].text = str(High)
        rowalignment = row[4].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if "Medium" in GraphAll[i]:
            Medium = int(GraphAll[i]["Medium"])
            row[5].text = str(Medium)
            rowalignment = row[5].paragraphs[0]
            rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[5].text = str(Medium)
        rowalignment = row[5].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if "Low" in GraphAll[i]:
            Low = int(GraphAll[i]["Low"])
            row[6].text = str(Low)
            rowalignment = row[6].paragraphs[0]
            rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[6].text = str(Low)
        rowalignment = row[6].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sum = (int(Critical) + int(High) + int(Medium) + int(Low))
        row[7].text = str(sum)
        rowalignment = row[7].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if numberoftable >= len(GraphAll):
        # A merge is specified
        numberoftable += 1
        row = table7.add_row().cells
        col0 = table7.cell(numberoftable, 0)
        col1 = table7.cell(numberoftable, 2)
        col0.merge(col1)
        row[0].text = "Total"
        rowalignment = row[0].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading_list = locals()
        shading_list['shading_elm_' + str(colsNum)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStrDefalt))
        table7.rows[len(GraphAll) + 1].cells[0]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(colsNum)])

        row[3].text = str(GraphTotol["Critical"])
        rowalignment = row[3].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        colorStr = '8139B8'
        shading_list['shading_elm_' + str(colsNum)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStr))
        table7.rows[len(GraphAll) + 1].cells[3]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(colsNum)])

        row[4].text = str(GraphTotol["High"])
        rowalignment = row[4].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        colorStr = 'FF0000'
        shading_list['shading_elm_' + str(colsNum)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStr))
        table7.rows[len(GraphAll) + 1].cells[4]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(colsNum)])

        row[5].text = str(GraphTotol["Medium"])
        rowalignment = row[5].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        colorStr = 'FFD400'
        shading_list['shading_elm_' + str(colsNum)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStr))
        table7.rows[len(GraphAll) + 1].cells[5]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(colsNum)])

        row[6].text = str(GraphTotol["Low"])
        rowalignment = row[6].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        colorStr = 'EFFF00'
        shading_list['shading_elm_' + str(colsNum)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStr))
        table7.rows[len(GraphAll) + 1].cells[6]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(colsNum)])

        row[7].text = str(GraphTotol["Critical"] + GraphTotol["High"] +
                          GraphTotol["Medium"] + GraphTotol["Low"])
        rowalignment = row[7].paragraphs[0]
        rowalignment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        colorStr = '0099FF'
        shading_list['shading_elm_' + str(colsNum)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStr))
        table7.rows[len(GraphAll) + 1].cells[7]._tc.get_or_add_tcPr().append(
            shading_list['shading_elm_' + str(colsNum)])

        for i in range(8):
            table7.rows[len(GraphAll) +
                        1].cells[i].paragraphs[0].runs[0].font.bold = True

    for cell in table7.columns[0].cells:
        cell.width = Inches(0.4)
    for cell in table7.columns[1].cells:
        cell.width = Inches(2)
    for cell in table7.columns[2].cells:
        cell.width = Inches(1.2)

    doc.add_page_break()

    doc.add_heading("5.3 Infrastructure Vulnerability Detail", 2)
    # doc.add_paragraph()

    id = 0
    for i in levelall:
        id += 1
        autoTable(id, i)
        doc.add_heading()
    filename = filename.split(".csv")
    doc.save('D:/github/WebReport/backend/uploads/'+filename[0]+".docx")
    # os.system('D:/github/WebReport/backend/api/templates/Vulnerability.docx')

    directory = "D:/github/WebReport/backend/uploads"

    files_in_directory = os.listdir(directory)
    filtered_files = [
        file for file in files_in_directory if file.endswith(".csv")]
    for file in filtered_files:
        path_to_file = os.path.join(directory, file)
        os.remove(path_to_file)

    directoryJson = "D:/github/WebReport/backend/api/sources"

    files_in_directory = os.listdir(directoryJson)
    filtered_files = [
        file for file in files_in_directory if file.endswith(".json")]
    for file in filtered_files:
        path_to_file = os.path.join(directoryJson, file)
        os.remove(path_to_file)

    return True
