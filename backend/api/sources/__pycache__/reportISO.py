from random import setstate
from docx import *
import os
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement, ns
import datetime
from matplotlib import cm
import pandas
import json
import csv
import matplotlib.pyplot as plt
from docx.oxml.ns import qn
from visions import Object

doc = Document()

# page margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(3)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Set page size A4
doc.sections[0].page_height = Cm(29.7)
doc.sections[0].page_width = Cm(21)

# Set Font Page default
style = doc.styles['Normal']
fontdefault = style.font
fontdefault.name = 'TH Sarabun New'
fontdefault.size = Pt(16)


def page1():
    # Header title
    header_title = doc.add_paragraph()
    header_title.alignment = 1
    # title.paragraph_format.space_after = Cm(1)
    header_title = header_title.add_run(
        "รายงานสรุปผลการตรวจสอบช่องโหว่ของระบบ Cloud Dell\nบริษัท อินเทอร์เน็ตประเทศไทย จำกัด(มหาชน) ประจำปี 2564")
    # paragraph.paragraph_format.space_before = Pt(3)
    font1 = header_title.style
    font1.font.size = Pt(16)

    Figure1 = doc.add_picture(
        'D:/github/WebReport/backend/api/image/horizental.jpg', width=Cm(17))
    Figure1 = doc.paragraphs[-1]
    Figure1.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Define image to right
    header_title.bold = True

    # Objective
    title = doc.add_heading(
        'วัตถุประสงค์', 1)
    title_style = title.style
    title_style.font.size = Pt(16)
    rFonts = title_style.element.rPr.rFonts
    rFonts.set(qn("w:hAnsiTheme"), 'TH Sarabun New')

    # content
    paragraph1 = doc.add_paragraph(
        "เพื่อค้นหาจุดอ่อนและช่องโหว่บนระบบโครงสร้างพื้นฐานเทคโนโลยีสารสนเทศของบริการ โดยใช้เทคนิคกาตรวจสอบแบบ Vulnerability Assessment อ้างอิงจากมาตรฐาน Common Vulnerability Exposure, CVE และ Common Vulnerability Scoring System, CVSS-SIG ",
        style='List Number 3')
    paragraph1.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    font = paragraph1.style
    font.font.size = Pt(14)

    paragraph2 = doc.add_paragraph(
        "เพื่อนำเสนอข้อมูลจุดอ่อนหรือช่องโหว่ที่พบจากการตรวจสอบบนระบบโครงสร้างพื้นฐานเทคโนโลยีสารสนเทศของบริการและเสนอแนะแนวทางในการแก้ไขจุดอ่อนและช่องโหว่ที่ตรวจพบ เพื่อให้ผู้ดูแลระบบเร่งดำเนินการ แก้ไข",
        style='List Number 3')
    paragraph2.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    font = paragraph2.style
    font.font.size = Pt(14)

    # target
    title = doc.add_heading(
        'เป้าหมาย', 1)
    title_style = title.style
    title_style.font.size = Pt(16)
    rFonts = title_style.element.rPr.rFonts
    rFonts.set(qn("w:hAnsiTheme"), 'TH Sarabun New')

    # content
    paragraph3 = doc.add_paragraph(
        "เพื่อให้ทราบถึงจุดอ่อนและช่องโหว่ของเครื่องที่ให้บริการ รวมถึงอุปกรณ์เครือข่ายภายในระบบสารสนเทศของ บริการ",
        style='List Number 3')
    paragraph3.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    font = paragraph3.style
    font.font.size = Pt(14)

    paragraph4 = doc.add_paragraph(
        "เพื่อนำเสนอข้อมูลจุดอ่อนหรือช่องโหว่ที่พบให้ผู้ดูแลระบบผู้ดูแลเครือข่ายและผู้ดูแลระบบงานรับทราบถึงช่อง โหว่ที่ตรวจพบเพื่อดำเนินการแก้ไขปรับปรุงให้ระบบมีความแข็งแกร่งและยกระดับความมั่นคงปลอดภัยระบบ สารสนเทศของบริการ",
        style='List Number 3')
    paragraph4.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    font = paragraph4.style
    font.font.size = Pt(14)

    # target
    title = doc.add_heading(
        'แนวทางการประเมินความเสี่ยงที่พบ', 1)
    title_style = title.style
    title_style.font.size = Pt(16)
    rFonts = title_style.element.rPr.rFonts
    rFonts.set(qn("w:hAnsiTheme"), 'TH Sarabun New')

    paragraph5 = doc.add_paragraph(
        "       การตรวจสอบจุดอ่อนและช่องโหว่ของเครื่องที่ให้บริการและอุปกรณ์เครือข่ายภายในโครงสร้างระบบสารสนเทศ    ดำเนินการตรวจสอบ โดยพิจารณาตามระดับความรุนแรงของจุดอ่อนและช่องโหว่ที่ตรวจพบ ซึ่งอ้างอิงตามมาตรฐาน CommonVulnerability Scoring System(CVSS)  ใช้เป็นมาตรฐานสากลโดยสามารถแบ่งระดับความรุนแรงของช่องโหว่ดังนี้")
    paragraph5.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    font = paragraph5.style
    font.font.size = Pt(14)

    paragraph6 = doc.add_paragraph(style='List Bullet 2')
    paragraph6.add_run('''ความรุนแรงระดับวิกฤติ (CRITICAL)'''
                       ).font.highlight_color = RGBColor(0x42, 0x24, 0xE9)
    paragraph6.add_run(' หมายถึงจุดอ่อนหรือช่องโหว่ที่มีความเสี่ยงต่อการถูกบุกรุกระบบระดับวิกฤติ ผู้บุกรุกระบบสามารถใช้ช่องโหว่ที่ตรวจพบนี้โจมตีระบบได้ทันที และสร้างความเสียหายต่อระบบสารสนเทศใน ระดับวิกฤติ')
    paragraph6.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

    paragraph7 = doc.add_paragraph(style='List Bullet 2')
    paragraph7.add_run('''ความรุนแรงระดับสูง (HIGH)'''
                       ).font.highlight_color = WD_COLOR_INDEX.VIOLET
    paragraph7.add_run(
        ' หมายถึงจุดอ่อนหรือช่องโหว่ที่มีความเสี่ยงต่อการถูกบุกรุกระบบระดับสูง ผู้บุกรุกระบบสามารถใช้ช่องโหว่ที่ตรวจพบนี้โจมตีระบบได้ทันที และสร้างความเสียหายต่อระบบสารสนเทศในระดับสูง')
    paragraph7.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

    paragraph8 = doc.add_paragraph(style='List Bullet 2')
    paragraph8.add_run('''ความรุนแรงระดับกลาง (MEDIUM)'''
                       ).font.highlight_color = WD_COLOR_INDEX.YELLOW
    paragraph8.add_run(
        ' หมายถึงจุดอ่อนหรือช่องโหว่ที่มีความเสี่ยงต่อการถูกบุกรุกระบบระดับกลาง และผลกระทบของการบุกรุกระบบจะทำให้ระบบสารสนเทศมีความเสียหายในระดับกลาง')
    paragraph8.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

    paragraph9 = doc.add_paragraph(style='List Bullet 2')
    paragraph9.add_run('''ความรุนแรงระดับต่ำ (LOW)'''
                       ).font.highlight_color = WD_COLOR_INDEX.GREEN
    paragraph9.add_run(
        ' หมายถึงจุดอ่อนหรือช่องโหว่ที่มีความเสี่ยงต่อการถูกบุกรุกระบบระดับต่ำ และผลกระทบของการบุกรุกระบบทำให้ระบบสารสนเทศมีความเสียหายในระดับต่ำ')
    paragraph9.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

    paragraph10 = doc.add_paragraph(style='List Bullet 2')
    paragraph10.add_run('''ไม่พบความรุนแรง (INFO)'''
                        ).font.highlight_color = WD_COLOR_INDEX.BLUE
    paragraph10.add_run(
        ' หมายถึงรายละเอียดทั่วไปของระบบสารสนเทศ ซึ่งไม่มีผลกระทบต่อความเสียหายของระบบสารสนเทศ')
    paragraph10.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY


page1()

# table = doc.add_table(rows=0, cols=2, style='Table Grid')
# row = table.add_row().cells
# paragraph = row[0].paragraphs[0]
# run = paragraph.add_run()
# run.add_picture('D:/github/WebReport/backend/api/image/horizental.jpg', width = 1400000, height = 1400000)


# save and open

doc.save("reportISO.docx")
os.system("reportISO.docx")
